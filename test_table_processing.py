#!/usr/bin/env python3
"""
Тестовый скрипт для проверки новой функциональности обработки таблиц
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from services.template_processor_service import TemplateProcessorService
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

def test_table_processing():
    """Тестируем новую функцию обработки таблиц"""
    
    # Создаем тестовый документ с таблицей
    doc = Document()
    
    # Добавляем параграф
    doc.add_paragraph("Договор подряда")
    
    # Добавляем таблицу
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Заполняем таблицу
    table.cell(0, 0).text = "ИСПОЛНИТЕЛЬ:"
    table.cell(0, 1).text = "ООО \"Тестовая компания\""
    
    table.cell(1, 0).text = "ИНН:"
    table.cell(1, 1).text = "1234567890"
    
    table.cell(2, 0).text = "ЗАКАЗЧИК:"
    table.cell(2, 1).text = "_________________"
    
    # Добавляем еще один параграф
    doc.add_paragraph("Подписи сторон")
    
    # Тестируем новую функцию напрямую (без инициализации Gemini)
    from services.template_processor_service import TemplateProcessorService
    
    # Создаем экземпляр сервиса, но не инициализируем Gemini
    service = TemplateProcessorService.__new__(TemplateProcessorService)
    service.prompt_manager = None  # Не нужен для теста
    
    map_for_gemini, coords_dictionary = service._index_runs_and_build_map(doc)
    
    print("=== РЕЗУЛЬТАТ ТЕСТИРОВАНИЯ ===")
    print(f"Количество run-ов: {len(coords_dictionary)}")
    print(f"Длина карты: {len(map_for_gemini)} символов")
    print("\n=== КАРТА ДОКУМЕНТА ===")
    print(map_for_gemini)
    print("\n=== КООРДИНАТЫ RUN-ОВ ===")
    for run_id, run_obj in coords_dictionary.items():
        print(f"{run_id}: '{run_obj.text}'")
    
    # Проверяем, что таблица представлена в Markdown формате
    if "|" in map_for_gemini and ":---" in map_for_gemini:
        print("\n✅ ТАБЛИЦА УСПЕШНО ПРЕОБРАЗОВАНА В MARKDOWN!")
    else:
        print("\n❌ ОШИБКА: Таблица не преобразована в Markdown формат")

if __name__ == "__main__":
    test_table_processing()
