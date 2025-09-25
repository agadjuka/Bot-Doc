"""
Document Generator Service for creating .docx files from templates
"""
import io
import os
from typing import Dict, Any, Optional
from docxtpl import DocxTemplate


class DocumentGeneratorService:
    """Service for generating .docx documents from templates"""
    
    def __init__(self, templates_dir: str = "templates"):
        """
        Initialize the Document Generator Service
        
        Args:
            templates_dir: Directory containing template files
        """
        self.templates_dir = templates_dir
        if not os.path.exists(templates_dir):
            os.makedirs(templates_dir)
            print(f"📁 Создана папка для шаблонов: {templates_dir}")
    
    def fill_document(self, template_path: str, data: Dict[str, Any]) -> bytes:
        """
        Fill a document template with provided data
        
        Args:
            template_path: Path to the .docx template file
            data: Dictionary with data to fill the template
            
        Returns:
            Bytes of the filled document
            
        Raises:
            FileNotFoundError: If template file doesn't exist
            Exception: If document generation fails
        """
        try:
            # Check if template file exists
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Шаблон не найден: {template_path}")
            
            # Load the template
            template = DocxTemplate(template_path)
            
            # Clean the data - replace None values with empty strings
            cleaned_data = {}
            for key, value in data.items():
                if value is None:
                    cleaned_data[key] = ""
                else:
                    cleaned_data[key] = str(value)
            
            # Render the template with data
            template.render(cleaned_data)
            
            # Save to memory
            doc_buffer = io.BytesIO()
            template.save(doc_buffer)
            doc_buffer.seek(0)
            
            print(f"✅ Документ успешно сгенерирован из шаблона: {template_path}")
            return doc_buffer.getvalue()
            
        except FileNotFoundError as e:
            print(f"❌ Ошибка: {e}")
            raise
        except Exception as e:
            print(f"❌ Ошибка при генерации документа: {e}")
            raise
    
    def get_template_path(self, template_name: str) -> str:
        """
        Get full path to a template file
        
        Args:
            template_name: Name of the template file (with or without .docx extension)
            
        Returns:
            Full path to the template file
        """
        if not template_name.endswith('.docx'):
            template_name += '.docx'
        
        return os.path.join(self.templates_dir, template_name)
    
    def list_available_templates(self) -> list:
        """
        List all available template files
        
        Returns:
            List of template file names
        """
        try:
            if not os.path.exists(self.templates_dir):
                return []
            
            templates = []
            for file in os.listdir(self.templates_dir):
                if file.endswith('.docx'):
                    templates.append(file)
            
            return templates
        except Exception as e:
            print(f"❌ Ошибка при получении списка шаблонов: {e}")
            return []
    
    def validate_template(self, template_path: str) -> bool:
        """
        Validate that a template file exists and is accessible
        
        Args:
            template_path: Path to the template file
            
        Returns:
            True if template is valid, False otherwise
        """
        try:
            if not os.path.exists(template_path):
                return False
            
            # Try to load the template
            template = DocxTemplate(template_path)
            return True
            
        except Exception as e:
            print(f"❌ Ошибка валидации шаблона {template_path}: {e}")
            return False
    
    def create_sample_template(self, template_name: str = "template_test.docx") -> str:
        """
        Create a sample template file for testing
        
        Args:
            template_name: Name of the template file to create
            
        Returns:
            Path to the created template file
        """
        try:
            from docx import Document
            
            template_path = self.get_template_path(template_name)
            
            # Create a new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Информация о компании', 0)
            
            # Add company information with placeholders
            doc.add_heading('Основная информация', level=1)
            
            p1 = doc.add_paragraph('Полное наименование: ')
            p1.add_run('{{COMPANY_FULL_NAME}}').bold = True
            
            p2 = doc.add_paragraph('Краткое наименование: ')
            p2.add_run('{{COMPANY_SHORT_NAME}}').bold = True
            
            p3 = doc.add_paragraph('ИНН: ')
            p3.add_run('{{INN}}').bold = True
            
            p4 = doc.add_paragraph('КПП: ')
            p4.add_run('{{KPP}}').bold = True
            
            p5 = doc.add_paragraph('ОГРН: ')
            p5.add_run('{{OGRN}}').bold = True
            
            doc.add_heading('Адрес', level=1)
            p6 = doc.add_paragraph('Юридический адрес: ')
            p6.add_run('{{LEGAL_ADDRESS}}').bold = True
            
            doc.add_heading('Руководство', level=1)
            p7 = doc.add_paragraph('Должность: ')
            p7.add_run('{{DIRECTOR_POSITION}}').bold = True
            
            p8 = doc.add_paragraph('ФИО: ')
            p8.add_run('{{DIRECTOR_FULL_NAME}}').bold = True
            
            # Save the document
            doc.save(template_path)
            print(f"✅ Создан тестовый шаблон: {template_path}")
            
            return template_path
            
        except ImportError:
            print("❌ Для создания шаблона требуется библиотека python-docx")
            print("Установите её командой: pip install python-docx")
            raise
        except Exception as e:
            print(f"❌ Ошибка при создании тестового шаблона: {e}")
            raise
