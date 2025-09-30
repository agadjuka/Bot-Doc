#!/usr/bin/env python3
"""
Простой тестер для проверки системы обработки шаблонов
"""
import json
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

# Загружаем переменные окружения
load_dotenv("env.local")

# Устанавливаем переменную окружения для Google Cloud
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "bot-doc-473208-706e6adceee1.json"

def load_gemini_response():
    """Загружает ответ от Gemini из файла"""
    try:
        with open("gemini_request_TemplateProcessorService_20250928_202156_318890_response.json", "r", encoding="utf-8") as f:
            response_data = json.load(f)
        
        gemini_response = response_data["response_text"]
        print(f"📄 Загружен ответ от Gemini ({len(gemini_response)} символов)")
        return gemini_response
        
    except Exception as e:
        print(f"❌ Ошибка загрузки ответа от Gemini: {e}")
        return None

def parse_gemini_response(gemini_response):
    """Парсит ответ от Gemini и извлекает план правок"""
    print("\n🔍 Парсим ответ от Gemini...")
    
    try:
        # Clean the response (remove markdown formatting if present)
        cleaned_response = gemini_response.strip()
        
        # Remove markdown code blocks
        if cleaned_response.startswith('```json'):
            cleaned_response = cleaned_response[7:]
        elif cleaned_response.startswith('```'):
            cleaned_response = cleaned_response[3:]
        
        if cleaned_response.endswith('```'):
            cleaned_response = cleaned_response[:-3]
        
        cleaned_response = cleaned_response.strip()
        
        # Try to find JSON array in the response
        json_start = cleaned_response.find('[')
        json_end = cleaned_response.rfind(']') + 1
        
        if json_start != -1 and json_end > json_start:
            json_text = cleaned_response[json_start:json_end]
            edits_plan = json.loads(json_text)
        else:
            edits_plan = json.loads(cleaned_response)
        
        print(f"✅ Извлечено {len(edits_plan)} правок:")
        for i, edit in enumerate(edits_plan):
            print(f"  {i+1}. run_id={edit['run_id']}, field_name='{edit['field_name']}'")
        
        return edits_plan
        
    except Exception as e:
        print(f"❌ Ошибка парсинга JSON: {e}")
        return []

def load_test_document():
    """Загружает тестовый Word-документ"""
    try:
        doc_path = "ДОГОВОР С АВТОКИТ!.docx"
        if not os.path.exists(doc_path):
            print(f"❌ Тестовый документ не найден: {doc_path}")
            return None
        
        with open(doc_path, "rb") as f:
            file_bytes = f.read()
        
        print(f"📄 Загружен тестовый документ: {len(file_bytes)} байт")
        return file_bytes
        
    except Exception as e:
        print(f"❌ Ошибка загрузки документа: {e}")
        return None

def index_runs_and_build_map(doc_object):
    """Создает карту run'ов для документа"""
    print("\n📊 Создаем карту run'ов...")
    
    try:
        map_for_gemini = ""
        coords_dictionary = {}
        run_counter = 0
        
        # Process document body elements in order (paragraphs and tables)
        for element in doc_object._body._body:
            if element.tag.endswith('p'):  # Paragraph
                # Process paragraph
                paragraph = None
                for p in doc_object.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                
                if paragraph:
                    for run in paragraph.runs:
                        # Generate unique ID for this run
                        run_id = f"run-{run_counter}"
                        
                        # Add run to coordinates dictionary
                        coords_dictionary[run_id] = run
                        
                        # Add run to map for Gemini
                        map_for_gemini += f"[{run_id}]{run.text}"
                        
                        # Increment counter
                        run_counter += 1
                    
                    # Add newline after each paragraph
                    map_for_gemini += "\n"
            
            elif element.tag.endswith('tbl'):  # Table
                # Process table
                table = None
                for t in doc_object.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table:
                    # Build matrix of runs for this table
                    table_matrix = []
                    
                    for row in table.rows:
                        row_runs = []
                        for cell in row.cells:
                            cell_runs = []
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    # Generate unique ID for this run
                                    run_id = f"run-{run_counter}"
                                    
                                    # Add run to coordinates dictionary
                                    coords_dictionary[run_id] = run
                                    
                                    # Add run to cell runs list
                                    cell_runs.append(f"[{run_id}]{run.text}")
                                    
                                    # Increment counter
                                    run_counter += 1
                            
                            # Join all runs in this cell with <br> for line breaks
                            cell_text = "<br>".join(cell_runs)
                            row_runs.append(cell_text)
                        
                        table_matrix.append(row_runs)
                    
                    # Generate Markdown table representation
                    if table_matrix:
                        # Add table header
                        if len(table_matrix) > 0:
                            header_row = "| " + " | ".join(table_matrix[0]) + " |"
                            map_for_gemini += header_row + "\n"
                            
                            # Add separator row
                            separator_row = "| " + " | ".join([":---"] * len(table_matrix[0])) + " |"
                            map_for_gemini += separator_row + "\n"
                            
                            # Add data rows
                            for row in table_matrix[1:]:
                                data_row = "| " + " | ".join(row) + " |"
                                map_for_gemini += data_row + "\n"
                        
                        # Add newline after table
                        map_for_gemini += "\n"
        
        print(f"📊 Найдено {len(coords_dictionary)} run'ов в документе")
        print(f"🔍 Первые 10 run_id: {list(coords_dictionary.keys())[:10]}")
        print(f"🔍 Последние 10 run_id: {list(coords_dictionary.keys())[-10:]}")
        
        return map_for_gemini, coords_dictionary
        
    except Exception as e:
        print(f"❌ Ошибка создания карты run'ов: {e}")
        return "", {}

def test_edits_application(doc, coords_dictionary, edits_plan):
    """Тестирует применение правок к документу"""
    print(f"\n🔧 Тестируем применение {len(edits_plan)} правок...")
    
    # Создаем копию документа для тестирования
    original_bytes = BytesIO()
    doc.save(original_bytes)
    original_bytes.seek(0)
    
    test_doc = Document(original_bytes)
    _, test_coords_dictionary = index_runs_and_build_map(test_doc)
    
    # Применяем правки
    successful_edits = 0
    failed_edits = 0
    
    for i, edit in enumerate(edits_plan):
        run_id = edit['run_id']
        field_name = edit['field_name']
        
        print(f"\n🔍 Правка {i+1}/{len(edits_plan)}: run_id={run_id}, field_name='{field_name}'")
        
        # Ищем run в тестовом документе
        test_run = test_coords_dictionary.get(run_id)
        
        if not test_run:
            print(f"❌ Run {run_id} не найден в тестовом документе")
            print(f"📊 Доступные run_id: {list(test_coords_dictionary.keys())[:10]}...")
            failed_edits += 1
            continue
        
        print(f"✅ Run {run_id} найден")
        print(f"📝 Текущий текст: '{test_run.text}'")
        
        # Применяем правку
        if field_name == "":
            print(f"🧹 Очищаю run {run_id}")
            test_run.text = ""
        else:
            print(f"✏️ Заменяю run {run_id} на '[{field_name}]'")
            test_run.text = f"[{field_name}]"
        
        print(f"📝 Новый текст: '{test_run.text}'")
        successful_edits += 1
    
    print(f"\n📊 Результат применения правок:")
    print(f"✅ Успешно применено: {successful_edits}")
    print(f"❌ Не удалось применить: {failed_edits}")
    
    return test_doc, successful_edits, failed_edits

def main():
    """Главная функция"""
    print("🚀 ЗАПУСК ПРОСТОГО ТЕСТА СИСТЕМЫ ОБРАБОТКИ ШАБЛОНОВ")
    print("=" * 60)
    
    # Шаг 1: Загружаем ответ от Gemini
    gemini_response = load_gemini_response()
    if not gemini_response:
        return
    
    # Шаг 2: Парсим ответ от Gemini
    edits_plan = parse_gemini_response(gemini_response)
    if not edits_plan:
        print("❌ Не удалось извлечь план правок")
        return
    
    # Шаг 3: Загружаем тестовый документ
    file_bytes = load_test_document()
    if not file_bytes:
        return
    
    # Шаг 4: Загружаем документ
    doc = Document(BytesIO(file_bytes))
    
    # Шаг 5: Создаем карту run'ов
    map_for_gemini, coords_dictionary = index_runs_and_build_map(doc)
    
    # Шаг 6: Тестируем применение правок
    test_doc, successful_edits, failed_edits = test_edits_application(
        doc, coords_dictionary, edits_plan
    )
    
    # Шаг 7: Сохраняем результаты
    try:
        test_output_path = "test_output_with_edits.docx"
        test_doc.save(test_output_path)
        print(f"💾 Тестовый документ сохранен: {test_output_path}")
    except Exception as e:
        print(f"❌ Ошибка сохранения: {e}")
    
    # Шаг 8: Итоговый отчет
    print("\n" + "=" * 60)
    print("📋 ИТОГОВЫЙ ОТЧЕТ")
    print("=" * 60)
    print(f"✅ Успешно применено правок: {successful_edits}")
    print(f"❌ Не удалось применить: {failed_edits}")
    print(f"📊 Процент успеха: {(successful_edits / len(edits_plan)) * 100:.1f}%")
    
    if failed_edits > 0:
        print("\n🔍 РЕКОМЕНДАЦИИ ПО УЛУЧШЕНИЮ:")
        print("1. Проверьте соответствие run_id в ответе Gemini и в документе")
        print("2. Убедитесь, что документ не был изменен после создания карты run'ов")
        print("3. Проверьте логику индексации run'ов")
    else:
        print("\n🎉 ВСЕ ПРАВКИ ПРИМЕНЕНЫ УСПЕШНО!")

if __name__ == "__main__":
    main()
