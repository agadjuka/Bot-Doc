"""
Template Processor Service for automatic field detection and marking in documents.
This service analyzes documents and identifies fields that need to be filled with company data.
"""

import json
import logging
import os
import re
import io
from typing import Dict, List, Tuple
from io import BytesIO

import google.generativeai as genai
from google.oauth2 import service_account
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import qn

from config.prompts import PromptManager

# ВРЕМЕННЫЙ ДЕБАГ - можно удалить после отладки
try:
    from debug_gemini_logger import log_gemini_request, log_gemini_response
    DEBUG_GEMINI = True
except ImportError:
    DEBUG_GEMINI = False
    def log_gemini_request(*args, **kwargs): return ""
    def log_gemini_response(*args, **kwargs): return ""

logger = logging.getLogger(__name__)


class TemplateProcessorService:
    """
    Service for processing document templates using simplified analysis strategy.
    Analyzes documents and creates two files: preview for user and smart template for storage.
    """
    
    def __init__(self):
        """
        Initialize the TemplateProcessorService using Google Cloud authentication.
        """
        self.prompt_manager = PromptManager()
        self._initialize_gemini()
        logger.info("TemplateProcessorService initialized successfully")
    
    def _remove_highlighting(self, run):
        """
        Remove yellow highlighting from a run by clearing background color.
        
        Args:
            run: python-docx run object
        """
        try:
            # Remove highlighting by clearing the highlight property
            if hasattr(run, '_element'):
                # Remove w:highlight attribute if it exists
                if run._element.get(qn('w:highlight')):
                    del run._element.attrib[qn('w:highlight')]
                
                # Remove w:shd (shading) attribute if it exists
                if run._element.get(qn('w:shd')):
                    del run._element.attrib[qn('w:shd')]
                
                # Also check for highlighting in the run's properties
                run_props = run._element.find(qn('w:rPr'))
                if run_props is not None:
                    highlight_elem = run_props.find(qn('w:highlight'))
                    if highlight_elem is not None:
                        run_props.remove(highlight_elem)
                    
                    shading_elem = run_props.find(qn('w:shd'))
                    if shading_elem is not None:
                        run_props.remove(shading_elem)
                
                # Additional cleanup for any background color formatting
                # Remove any w:color with background color
                color_elem = run_props.find(qn('w:color')) if run_props is not None else None
                if color_elem is not None:
                    # Check if it's a background color (usually has w:val="auto" or specific color)
                    if color_elem.get(qn('w:val')) in ['auto', 'yellow', 'FFFF00']:
                        run_props.remove(color_elem)
            
            # Also try to remove highlighting using python-docx properties
            if hasattr(run, 'font') and hasattr(run.font, 'highlight_color'):
                try:
                    run.font.highlight_color = None
                except:
                    pass
            
        except Exception as e:
            # Не критично, продолжаем работу
            pass

    def _initialize_gemini(self):
        """Initialize Google Gemini AI service using Google Cloud authentication"""
        try:
            # Get credentials file path
            credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
            if not credentials_path:
                logger.error("GOOGLE_APPLICATION_CREDENTIALS not set")
                raise ValueError("GOOGLE_APPLICATION_CREDENTIALS is required")
            
            if not os.path.exists(credentials_path):
                logger.error(f"Credentials file not found: {credentials_path}")
                raise FileNotFoundError(f"Credentials file not found: {credentials_path}")
            
            # Load service account credentials
            credentials = service_account.Credentials.from_service_account_file(
                credentials_path,
                scopes=['https://www.googleapis.com/auth/generative-language']
            )
            
            # Configure Gemini API with service account credentials
            genai.configure(credentials=credentials)
            
            # Initialize model (using flash model for better performance)
            self.model = genai.GenerativeModel('gemini-2.5-flash')
            logger.info("Gemini AI service initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize Gemini AI service: {e}")
            raise
    
    def _create_document_json_map(self, doc_object: Document) -> Tuple[Dict, Dict[str, any]]:
        """
        Create structured JSON representation of document for Gemini analysis.
        Creates hierarchical structure that mirrors document structure.
        
        Args:
            doc_object: python-docx Document object
            
        Returns:
            Tuple of (document_json, coords_dictionary)
            - document_json: Hierarchical JSON structure representing document
            - coords_dictionary: Dictionary mapping run_id to run objects
        """
        try:
            # Initialize variables
            document_json = {'body': []}
            coords_dictionary = {}
            run_counter = 0
            
            def process_paragraph(paragraph):
                """Process a paragraph and return its JSON representation"""
                nonlocal run_counter
                paragraph_data = {'type': 'paragraph', 'runs': []}
                
                # Process all runs in the paragraph
                for run in paragraph.runs:
                    # Generate unique ID for this run
                    run_id = f"run-{run_counter}"
                    
                    # Add run to coordinates dictionary
                    coords_dictionary[run_id] = run
                    
                    # Add run to paragraph data
                    paragraph_data['runs'].append({
                        'id': run_id,
                        'text': run.text
                    })
                    
                    # Increment counter
                    run_counter += 1
                
                # If paragraph has no runs (empty paragraph), create a dummy run
                if not paragraph.runs:
                    # Generate unique ID for dummy run
                    run_id = f"run-{run_counter}"
                    
                    # Create a dummy run by adding a run to the paragraph
                    dummy_run = paragraph.add_run("")
                    coords_dictionary[run_id] = dummy_run
                    
                    # Add dummy run to paragraph data
                    paragraph_data['runs'].append({
                        'id': run_id,
                        'text': ""
                    })
                    
                    # Increment counter
                    run_counter += 1
                
                return paragraph_data
            
            def process_cell(cell):
                """Process a table cell and return its JSON representation"""
                cell_data = {'type': 'cell', 'content': []}
                
                for paragraph in cell.paragraphs:
                    paragraph_data = process_paragraph(paragraph)
                    cell_data['content'].append(paragraph_data)
                
                return cell_data
            
            def process_table(table):
                """Process a table and return its JSON representation"""
                table_data = {'type': 'table', 'rows': []}
                
                for row in table.rows:
                    row_data = {'type': 'row', 'cells': []}
                    
                    for cell in row.cells:
                        cell_data = process_cell(cell)
                        row_data['cells'].append(cell_data)
                    
                    table_data['rows'].append(row_data)
                
                return table_data
            
            # Process document body elements in order
            for element in doc_object._body._body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find corresponding paragraph object
                    paragraph = None
                    for p in doc_object.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    
                    if paragraph:
                        paragraph_data = process_paragraph(paragraph)
                        document_json['body'].append(paragraph_data)
                
                elif element.tag.endswith('tbl'):  # Table
                    # Find corresponding table object
                    table = None
                    for t in doc_object.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        table_data = process_table(table)
                        document_json['body'].append(table_data)
            
            print(f"📊 JSON индексация: {len(coords_dictionary)} run-ов")
            print(f"🔍 Структура документа: {len(document_json['body'])} элементов в body")
            
            return document_json, coords_dictionary
            
        except Exception as e:
            print(f"❌ Ошибка JSON индексации: {e}")
            logger.error(f"Error creating document JSON map: {e}")
            return {'body': []}, {}
    
    def _apply_edits_to_runs(self, doc_object: Document, edits_plan: List[Dict[str, any]], coords_dictionary: Dict[str, any]) -> Tuple[bytes, bytes]:
        """
        Apply edits to documents using surgical approach based on new plan format.
        
        Args:
            doc_object: Original Document object
            edits_plan: List of edit dictionaries with target_run_ids, context_run_ids and field_name
            coords_dictionary: Dictionary mapping run_id to run objects
            
        Returns:
            Tuple of (preview_bytes, smart_template_bytes)
        """
        try:
            print(f"🔧 Применяю {len(edits_plan)} правок...")
            
            # Step 1: Create completely independent copies of the original document
            # Сохраняем оригинальный документ в байты и загружаем заново для каждой копии
            original_bytes = BytesIO()
            doc_object.save(original_bytes)
            original_bytes.seek(0)
            
            # Создаем preview документ из байтов
            preview_doc = Document(original_bytes)
            original_bytes.seek(0)
            
            # Создаем smart template документ из байтов
            smart_template_doc = Document(original_bytes)
            
            # Step 2: Rebuild coordinates dictionary for both copies
            _, preview_coords_dictionary = self._create_document_json_map(preview_doc)
            _, smart_template_coords_dictionary = self._create_document_json_map(smart_template_doc)
            
            # Step 3: Apply edits to both documents using new surgical algorithm
            for i, edit in enumerate(edits_plan):
                target_run_ids = edit.get('target_run_ids', [])
                context_run_ids = edit.get('context_run_ids', [])
                field_name = edit.get('field_name', '')
                
                print(f"🔍 Обрабатываю правку {i+1}/{len(edits_plan)}:")
                print(f"   target_run_ids={target_run_ids}")
                print(f"   context_run_ids={context_run_ids}")
                print(f"   field_name='{field_name}'")
                
                # Первый проход (Зачистка): Очищаем все target_run_ids
                print(f"🧹 Первый проход: зачистка {len(target_run_ids)} целевых run-ов")
                for run_id in target_run_ids:
                    preview_run = preview_coords_dictionary.get(run_id)
                    smart_template_run = smart_template_coords_dictionary.get(run_id)
                    
                    if preview_run:
                        print(f"   Очищаю preview run {run_id}: '{preview_run.text}' -> ''")
                        preview_run.text = ''
                    else:
                        print(f"   ❌ Run {run_id} не найден в preview документе")
                    
                    if smart_template_run:
                        print(f"   Очищаю smart_template run {run_id}: '{smart_template_run.text}' -> ''")
                        smart_template_run.text = ''
                    else:
                        print(f"   ❌ Run {run_id} не найден в smart_template документе")
                
                # Второй проход (Вставка и стилизация): Вставляем маркер в первый target_run
                if target_run_ids and field_name:
                    first_target_run_id = target_run_ids[0]
                    preview_run = preview_coords_dictionary.get(first_target_run_id)
                    smart_template_run = smart_template_coords_dictionary.get(first_target_run_id)
                    
                    if preview_run:
                        print(f"✏️ Второй проход: вставляю маркер в preview run {first_target_run_id}")
                        preview_run.text = f"[{field_name}]"
                        # Убираем подсветку и применяем красный жирный стиль
                        self._remove_highlighting(preview_run)
                        # Проверяем, что это действительно Run объект, а не Paragraph
                        if hasattr(preview_run, 'font'):
                            preview_run.font.color.rgb = RGBColor(255, 0, 0)
                            preview_run.bold = True
                        print(f"   Результат: '{preview_run.text}' (красный, жирный)")
                    else:
                        print(f"   ❌ Первый target run {first_target_run_id} не найден в preview")
                    
                    if smart_template_run:
                        print(f"✏️ Второй проход: вставляю маркер в smart_template run {first_target_run_id}")
                        smart_template_run.text = f"{{{{{field_name}}}}}"
                        # Убираем подсветку
                        self._remove_highlighting(smart_template_run)
                        print(f"   Результат: '{smart_template_run.text}'")
                    else:
                        print(f"   ❌ Первый target run {first_target_run_id} не найден в smart_template")
                
                # Третий проход (Очистка от желтизны): Убираем подсветку с context_run_ids
                print(f"🧽 Третий проход: очистка от желтизны {len(context_run_ids)} контекстных run-ов")
                for run_id in context_run_ids:
                    preview_run = preview_coords_dictionary.get(run_id)
                    
                    if preview_run:
                        print(f"   Убираю подсветку с preview run {run_id}: '{preview_run.text}'")
                        self._remove_highlighting(preview_run)
                    else:
                        print(f"   ❌ Context run {run_id} не найден в preview документе")
                
                print(f"✅ Правка {i+1} применена успешно")
            
            # Step 4: Save both documents to bytes
            # Save preview document
            preview_stream = BytesIO()
            preview_doc.save(preview_stream)
            preview_bytes = preview_stream.getvalue()
            
            # Save smart template document
            smart_template_stream = BytesIO()
            smart_template_doc.save(smart_template_stream)
            smart_template_bytes = smart_template_stream.getvalue()
            
            print(f"✅ Правки применены: preview {len(preview_bytes)} байт, template {len(smart_template_bytes)} байт")
            return preview_bytes, smart_template_bytes
            
        except Exception as e:
            print(f"❌ Ошибка при применении правок: {e}")
            logger.error(f"Error applying edits to runs: {e}")
            return b'', b''

    
    def _parse_gemini_edits_plan(self, gemini_response: str) -> List[Dict[str, any]]:
        """
        Parse JSON response from Gemini to extract edits plan.
        
        Args:
            gemini_response: Raw JSON response from Gemini
            
        Returns:
            List of edit dictionaries with target_run_ids, context_run_ids and field_name
        """
        try:
            # Clean the response (remove markdown formatting if present)
            cleaned_response = gemini_response.strip()
            
            # Remove markdown code blocks
            if cleaned_response.startswith('```json'):
                cleaned_response = cleaned_response[7:]
            elif cleaned_response.startswith('```'):
                cleaned_response = cleaned_response[3:]
            
            if cleaned_response.endswith('```'):
                cleaned_response = cleaned_response[:-3]
            
            cleaned_response = cleaned_response.strip()
            
            # Try multiple parsing strategies
            edits_plan = None
            
            # Strategy 1: Try to find JSON array in the response
            json_start = cleaned_response.find('[')
            json_end = cleaned_response.rfind(']') + 1
            
            if json_start != -1 and json_end > json_start:
                json_text = cleaned_response[json_start:json_end]
                try:
                    edits_plan = json.loads(json_text)
                except json.JSONDecodeError:
                    edits_plan = None
            
            # Strategy 2: Try to parse the whole response
            if edits_plan is None:
                try:
                    edits_plan = json.loads(cleaned_response)
                except json.JSONDecodeError:
                    edits_plan = None
            
            # Strategy 3: Try to extract JSON using regex
            if edits_plan is None:
                json_pattern = r'\[.*?\]'
                json_matches = re.findall(json_pattern, cleaned_response, re.DOTALL)
                if json_matches:
                    for json_match in json_matches:
                        try:
                            edits_plan = json.loads(json_match)
                            break
                        except json.JSONDecodeError:
                            continue
            
            if edits_plan is None:
                print(f"❌ Не удалось распарсить JSON из ответа")
                logger.error("Could not parse JSON from Gemini response")
                return []
            
            if not isinstance(edits_plan, list):
                logger.error("Gemini response is not a list")
                return []
            
            # Validate that each item has required fields for new format
            valid_edits = []
            for i, item in enumerate(edits_plan):
                if isinstance(item, dict):
                    # Check for new format fields
                    if 'target_run_ids' in item and 'context_run_ids' in item and 'field_name' in item:
                        # Validate that both run_ids fields are lists
                        if (isinstance(item['target_run_ids'], list) and 
                            isinstance(item['context_run_ids'], list)):
                            valid_edits.append(item)
                            print(f"📝 Правка {len(valid_edits)} (новый формат):")
                            print(f"   target_run_ids={item['target_run_ids']}")
                            print(f"   context_run_ids={item['context_run_ids']}")
                            print(f"   field_name='{item['field_name']}'")
                        else:
                            print(f"⚠️ target_run_ids и context_run_ids должны быть списками в правке {i+1}: {item}")
                    # Check for old format fields (backward compatibility)
                    elif 'run_ids' in item and 'field_name' in item:
                        if isinstance(item['run_ids'], list):
                            # Convert old format to new format
                            converted_item = {
                                'target_run_ids': item['run_ids'],
                                'context_run_ids': [],  # Empty context for old format
                                'field_name': item['field_name']
                            }
                            valid_edits.append(converted_item)
                            print(f"📝 Правка {len(valid_edits)} (старый формат, конвертирован):")
                            print(f"   target_run_ids={converted_item['target_run_ids']}")
                            print(f"   context_run_ids={converted_item['context_run_ids']}")
                            print(f"   field_name='{converted_item['field_name']}'")
                        else:
                            print(f"⚠️ run_ids должен быть списком в правке {i+1}: {item}")
                    else:
                        print(f"⚠️ Неверный формат правки {i+1}: {item}")
                        print(f"   Ожидаются поля: target_run_ids, context_run_ids, field_name (новый формат)")
                        print(f"   или: run_ids, field_name (старый формат)")
            
            print(f"✅ Извлечено {len(valid_edits)} валидных правок из {len(edits_plan)} элементов")
            logger.info(f"Successfully parsed {len(valid_edits)} valid edits from Gemini response")
            return valid_edits
            
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing JSON response from Gemini: {e}")
            print(f"❌ Ошибка парсинга JSON: {e}")
            return []
        except Exception as e:
            logger.error(f"Unexpected error parsing Gemini response: {e}")
            print(f"❌ Неожиданная ошибка: {e}")
            return []
    
    async def analyze_and_prepare_templates(self, file_bytes: bytes, file_format: str = '.docx') -> Tuple[bytes, bytes]:
        """
        Analyze document and prepare two files: preview for user and smart template for storage.
        
        Args:
            file_bytes: Document content as bytes
            file_format: File format ('.docx' or '.doc')
            
        Returns:
            Tuple of (preview_bytes, smart_template_bytes)
        """
        try:
            print(f"📄 Анализ документа: {len(file_bytes)} байт")
            
            # Step 1: Load document using python-docx for precise run-level analysis
            if file_format == '.docx':
                doc_object = Document(io.BytesIO(file_bytes))
            elif file_format == '.doc':
                print(f"❌ DOC формат не поддерживается")
                return b'', b''
            else:
                print(f"❌ Неподдерживаемый формат: {file_format}")
                return b'', b''
            
            # Step 2: Create structured JSON representation
            document_json, coords_dictionary = self._create_document_json_map(doc_object)
            
            if not document_json.get('body'):
                print(f"⚠️ Документ пустой")
                logger.warning("Document appears to be empty or could not be indexed")
                return b'', b''
            
            # Debug: Log first 500 characters of JSON and coords dictionary size
            json_str = json.dumps(document_json, indent=2, ensure_ascii=False)
            print(f"🔍 JSON структура (первые 500 символов):")
            print(json_str[:500] + "..." if len(json_str) > 500 else json_str)
            print(f"📊 Количество элементов в coords_dictionary: {len(coords_dictionary)}")
            
            # Convert JSON to text for Gemini (temporary approach)
            map_for_gemini = json_str
            
            # Step 3: Call Gemini for document analysis
            print(f"🤖 Отправляю в Gemini...")
            prompt = self.prompt_manager.get_document_analysis_prompt(map_for_gemini)
            
            # Send request to Gemini
            gemini_response = await self._send_gemini_request(prompt)
            
            if not gemini_response:
                print(f"❌ Пустой ответ от Gemini")
                logger.error("Empty response from Gemini")
                return b'', b''
            
            print(f"✅ Получен ответ от Gemini: {len(gemini_response)} символов")
            
            # Step 4: Parse Gemini response to get edits plan
            edits_plan = self._parse_gemini_edits_plan(gemini_response)
            
            if not edits_plan:
                print(f"❌ Не удалось получить план правок")
                logger.error("Failed to parse edits plan from Gemini response")
                return b'', b''
            
            # Step 5: Apply edits to documents
            preview_bytes, smart_template_bytes = self._apply_edits_to_runs(doc_object, edits_plan, coords_dictionary)
            
            if not preview_bytes or not smart_template_bytes:
                print(f"❌ Ошибка при применении правок")
                logger.error("Failed to apply edits to documents")
                return b'', b''
            
            return preview_bytes, smart_template_bytes
            
        except Exception as e:
            print(f"❌ Ошибка при анализе: {e}")
            logger.error(f"Error in document analysis: {e}")
            return b'', b''
    
    
    
    
    async def _send_gemini_request(self, prompt: str) -> str:
        """
        Send request to Gemini API.
        
        Args:
            prompt: Prompt to send to Gemini
            
        Returns:
            Response from Gemini
        """
        try:
            # ВРЕМЕННЫЙ ДЕБАГ - логируем запрос
            request_file = ""
            if DEBUG_GEMINI:
                request_file = log_gemini_request(
                    prompt=prompt,
                    service_name="TemplateProcessorService",
                    user_id=None,
                    additional_info={"prompt_type": "document_analysis"}
                )
            
            # Generate content using Gemini
            response = self.model.generate_content(prompt)
            
            if response.text:
                print(f"✅ Получен ответ от Gemini: {len(response.text)} символов")
                logger.info("Received response from Gemini")
                
                # ВРЕМЕННЫЙ ДЕБАГ - логируем ответ
                if DEBUG_GEMINI and request_file:
                    log_gemini_response(
                        response=response.text,
                        request_filepath=request_file,
                        success=True
                    )
                
                return response.text
            else:
                print(f"❌ Пустой ответ от Gemini")
                logger.error("Empty response from Gemini")
                
                # ВРЕМЕННЫЙ ДЕБАГ - логируем пустой ответ
                if DEBUG_GEMINI and request_file:
                    log_gemini_response(
                        response="",
                        request_filepath=request_file,
                        success=False,
                        error_message="Empty response from Gemini"
                    )
                
                return ""
                
        except Exception as e:
            print(f"❌ Ошибка Gemini: {e}")
            logger.error(f"Error sending request to Gemini: {e}")
            
            # ВРЕМЕННЫЙ ДЕБАГ - логируем ошибку
            if DEBUG_GEMINI and request_file:
                log_gemini_response(
                    response="",
                    request_filepath=request_file,
                    success=False,
                    error_message=str(e)
                )
            
            return ""
    
    # Old methods removed - using new surgical approach with edits plan